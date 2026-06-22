#!/usr/bin/env python
"""按指定模板改写源文件 → 输出改写后的 docx（可行性验证脚本）。

独立于飞书链路，仅复用 app 内的文档解析与 LLM provider。

用法：
  python scripts/rewrite_by_template.py \
      --source 源文件.docx --template 模板.docx [--output 改写后.docx]

效果说明：
- 以「模板」为基底文档（保留其主题字体、页面设置、页眉页脚），清空正文后按模板
  的章节骨架重新写入内容。
- 内容只来自「源文件」：缺少某节信息时显式写「（源文件中未提供相应内容）」，
  不臆造、不篡改源文件中的数值/名称/日期。
"""

from __future__ import annotations

import argparse
import asyncio
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "backend"))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.text.paragraph import Paragraph  # noqa: E402

from app.core.config import get_settings  # noqa: E402
from app.services.document_parser import parse_document  # noqa: E402
from app.services.llm_provider import (  # noqa: E402
    LLMConfig,
    LLMProvider,
    TokenUsage,
    build_chat_provider,
)

# 喂给「填充模板」单次调用的源文件内容字符预算；超出则先分片压缩成事实笔记再填充。
SOURCE_CHAR_BUDGET = 36_000
# 压缩阶段单片字符数。
CONDENSE_CHUNK_CHARS = 14_000

REWRITE_SYSTEM_PROMPT = """你是专业文档改写助手。任务：把「源文件内容」按「目标模板结构」\
重新组织、改写成一篇成稿。

关于模板结构的约定：
- 以 #、##、### 给出的是模板的章节标题与层级（每次模板都不同，必须照搬，不要套用任何固定模板）。
- 以「> 填写要求：」开头的行是该节的写作指引，告诉你这一节该写什么；它不是正文内容，
  不要原样抄进成稿，而要按其要求、用源文件里的事实去填写。

必须严格遵守：
1. 严格遵循模板给出的章节标题与层级，不增删、不重命名、不调整顺序。
2. 每一节的内容只能来自源文件，不得臆造、补全或编造源文件中不存在的信息。
3. 必须忠实保留源文件中的关键数据、数值、单位、名称、日期，不得篡改或四舍五入。
4. 若源文件缺少某节所需信息，该节正文只写「（源文件中未提供相应内容）」，不要用常识或外部知识补写。
5. 语言精炼、书面化，做合理的归纳与组织，但不改变事实。

输出格式（Markdown）：
- 模板的章节标题用对应级别的 #、##、### 等，与输入层级一致。
- 正文用普通段落；并列要点用 - 列表；数据对照用 Markdown 表格。
- 只输出改写后的成稿正文本身，不要任何解释、前言或结束语。
"""

CONDENSE_SYSTEM_PROMPT = """你是资料整理助手。请把给定的材料片段压缩成保真的事实笔记，\
用于后续按模板改写。要求：保留所有关键数据、数值、单位、名称、日期、结论；不得臆造、\
不得补全原文没有的信息；输出简洁的中文要点（可用 - 列表），不做评价、不加建议。"""


def _style_heading_level(paragraph: Paragraph) -> int | None:
    """从段落「样式名」推断标题级别；非标题样式返回 None。

    兼容内置英文样式（Heading 1 / Title）与本地化中文样式（标题 1 / 标题）。
    """
    style = paragraph.style
    name = (style.name if style is not None else "") or ""
    lowered = name.lower()
    if lowered.startswith("heading") or name.startswith("标题"):
        match = re.search(r"(\d+)", name)
        return int(match.group(1)) if match else 1
    if lowered == "title" or name == "题目":
        return 1
    return None


def _para_signals(paragraph: Paragraph, child) -> tuple[float | None, bool, bool]:
    """提取段落的排版信号：最大字号、是否（含）加粗、是否带列表编号。

    很多模板不用 Word 标题样式，而是靠「字号 + 加粗 + 编号」表达层级，这里据此兜底识别。
    """
    sizes = [r.font.size.pt for r in paragraph.runs if r.font.size is not None]
    max_size = max(sizes) if sizes else None
    visible_runs = [r for r in paragraph.runs if r.text.strip()]
    any_bold = any(bool(r.bold) for r in visible_runs)
    pPr = child.find(qn("w:pPr"))
    numbered = pPr is not None and pPr.find(qn("w:numPr")) is not None
    return max_size, any_bold, numbered


def _infer_level(paragraph: Paragraph, child) -> int | None:
    """综合样式名与排版信号推断层级；返回 None 表示「这是填写要求/正文」而非标题。

    优先级：① 标题样式名；② 大字号（>=14）→ 章级(1)；③ 加粗的编号小节 → 节级(2)；
    其余（不加粗的编号项、普通段落）视为该节的填写要求，不作标题。
    """
    level = _style_heading_level(paragraph)
    if level is not None:
        return level
    max_size, any_bold, _ = _para_signals(paragraph, child)
    if max_size is not None and max_size >= 14:
        return 1
    if any_bold:
        return 2
    return None


def extract_template_outline(template_path: Path) -> str:
    """把模板 docx 序列化成带层级的大纲文本，供模型理解目标结构与每节的写作要求。

    标题/小节 → 对应级别的 Markdown 标题；不加粗的编号项、普通段落作为「填写要求」带上
    （模板里常含「需列明政策名称…」「包括世界、中国市场规模…」这类指引）；表格记为占位提示。
    """
    document = Document(str(template_path))
    body = document.element.body
    lines: list[str] = []
    for child in body.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if not text:
                continue
            level = _infer_level(paragraph, child)
            if level is not None:
                lines.append("\n" + "#" * min(level, 6) + " " + text)
            else:
                lines.append(f"> 填写要求：{text}")
        elif tag.endswith("}tbl"):
            lines.append("> （此处模板含一个表格，若有对应数据请填为 Markdown 表格）")
    outline = "\n".join(lines).strip()
    return outline or "（模板未检测到明确的标题结构，请按源文件内容组织成一篇通顺成稿）"


def _chunk_text(text: str, chunk_chars: int) -> list[str]:
    paragraphs = [p for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    for para in paragraphs:
        if current and current_len + len(para) > chunk_chars:
            chunks.append("\n\n".join(current))
            current = []
            current_len = 0
        current.append(para)
        current_len += len(para)
    if current:
        chunks.append("\n\n".join(current))
    return chunks


async def condense_source(
    text: str,
    provider: LLMProvider,
    budget: int,
) -> tuple[str, TokenUsage]:
    """源文件过长时先分片压缩成保真事实笔记，控制在预算内再用于填充模板。"""
    usage = TokenUsage()
    if len(text) <= budget:
        return text, usage

    chunks = _chunk_text(text, CONDENSE_CHUNK_CHARS)
    print(f"  源文件较长（{len(text):,} 字符），分 {len(chunks)} 片压缩成事实笔记...")
    notes: list[str] = []
    for index, chunk in enumerate(chunks, start=1):
        result = await provider.complete(
            [
                {"role": "system", "content": CONDENSE_SYSTEM_PROMPT},
                {"role": "user", "content": f"材料片段 {index}/{len(chunks)}：\n{chunk}"},
            ]
        )
        notes.append(result.content.strip())
        usage += result.usage
        print(f"    片段 {index}/{len(chunks)} 压缩完成。")
    combined = "\n\n".join(notes)
    # 压缩后仍超预算则按预算硬截断（验证脚本够用），并提示。
    if len(combined) > budget:
        print(f"  压缩后仍有 {len(combined):,} 字符，截断至 {budget:,} 字符用于改写。")
        combined = combined[:budget]
    return combined, usage


async def rewrite_to_markdown(
    outline: str,
    source_text: str,
    provider: LLMProvider,
) -> tuple[str, TokenUsage]:
    user_prompt = (
        f"目标模板结构：\n{outline}\n\n"
        f"源文件内容：\n{source_text}\n\n"
        "请严格按上述模板结构，用源文件内容改写输出 Markdown 成稿。"
    )
    result = await provider.complete_until_done(
        [
            {"role": "system", "content": REWRITE_SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ]
    )
    return result.content, result.usage


# ----------------------- Markdown → docx 渲染 -----------------------


def _clear_body(document: Document) -> None:
    """清空模板正文，但保留 sectPr（页面设置/页眉页脚随之保留）。"""
    body = document.element.body
    for child in list(body.iterchildren()):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _add_runs_with_bold(paragraph: Paragraph, text: str) -> None:
    """把 **加粗** 片段渲染成粗体 run，其余为普通 run。"""
    for index, segment in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not segment:
            continue
        run = paragraph.add_run(segment)
        if index % 2 == 1:  # 位于 ** ** 之间的片段
            run.bold = True


def _paragraph_safe(document: Document, style_candidates: list[str]):
    """按候选顺序套用段落样式，模板缺该样式则跳过；都没有则返回无样式段落。

    返回 (paragraph, 实际使用的样式名 or None)。不同模板内置样式不一（如缺 List Bullet），
    这里逐个降级避免 KeyError 直接崩溃。
    """
    for style in style_candidates:
        try:
            return document.add_paragraph(style=style), style
        except KeyError:
            continue
    return document.add_paragraph(), None


def _add_heading_safe(document: Document, text: str, level: int):
    """套用 Heading 样式；模板缺对应样式时退化为加粗段落。"""
    try:
        return document.add_heading(text, level=min(level, 9))
    except KeyError:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        return paragraph


def _add_list_item(document: Document, text: str, *, ordered: bool, index: int) -> None:
    """渲染列表项；无专用列表样式时退化为带项目符号/序号前缀的段落。"""
    candidates = ["List Number"] if ordered else ["List Bullet"]
    paragraph, used = _paragraph_safe(document, candidates)
    if used is None:
        # 退化为普通段落（或 List Paragraph 缩进），自己补符号，保证可读。
        paragraph, used = _paragraph_safe(document, ["List Paragraph"])
        prefix = f"{index}. " if ordered else "• "
        _add_runs_with_bold(paragraph, prefix + text)
        return
    _add_runs_with_bold(paragraph, text)


def _is_table_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?[\s:\-\|]+\|?\s*", line)) and "-" in line


def _split_table_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def _render_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = row[c_idx] if c_idx < len(row) else ""
            if r_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def render_markdown_to_docx(
    markdown: str,
    template_path: Path,
    output_path: Path,
) -> None:
    """以模板为基底渲染 Markdown 成稿为 docx（继承模板主题/页眉页脚）。"""
    document = Document(str(template_path))
    _clear_body(document)

    lines = markdown.replace("\r\n", "\n").split("\n")
    i = 0
    n = len(lines)
    ordered_index = 0  # 连续有序列表的序号；遇到非有序项时归零。
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 表格：连续以 | 开头的行
        if stripped.startswith("|"):
            ordered_index = 0
            table_lines: list[str] = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [
                _split_table_row(tl)
                for tl in table_lines
                if not _is_table_separator(tl)
            ]
            _render_table(document, rows)
            continue

        # 标题
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            ordered_index = 0
            level = len(heading_match.group(1))
            _add_heading_safe(document, heading_match.group(2).strip(), level)
            i += 1
            continue

        # 无序列表
        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            ordered_index = 0
            _add_list_item(document, bullet_match.group(1).strip(), ordered=False, index=0)
            i += 1
            continue

        # 有序列表
        ordered_match = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if ordered_match:
            ordered_index += 1
            _add_list_item(
                document, ordered_match.group(1).strip(), ordered=True, index=ordered_index
            )
            i += 1
            continue

        # 普通段落
        ordered_index = 0
        paragraph = document.add_paragraph()
        _add_runs_with_bold(paragraph, stripped)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


# ----------------------- 主流程 -----------------------


async def run(source: Path, template: Path, output: Path) -> None:
    settings = get_settings()
    if not settings.llm_api_key:
        raise SystemExit("未配置 LLM_API_KEY，无法改写。请检查 .env。")

    provider = build_chat_provider(
        LLMConfig(
            provider=settings.llm_provider,
            base_url=settings.llm_base_url,
            api_key=settings.llm_api_key,
            model=settings.llm_model,
            timeout_ms=settings.llm_timeout_ms,
            max_tokens=settings.llm_max_tokens,
            temperature=settings.llm_temperature,
        ),
        settings.fallback_models,
    )

    print(f"[1/4] 解析模板结构：{template.name}")
    outline = extract_template_outline(template)
    print(f"      模板大纲：\n{_indent(outline)}\n")

    print(f"[2/4] 解析源文件：{source.name}")
    parsed_source = parse_document(source)
    print(
        f"      源文件类型 {parsed_source.file_type}，"
        f"约 {len(parsed_source.text):,} 字符。"
    )

    source_text, condense_usage = await condense_source(
        parsed_source.text, provider, SOURCE_CHAR_BUDGET
    )

    print(f"[3/4] 调用模型 {settings.llm_model} 按模板改写...")
    markdown, rewrite_usage = await rewrite_to_markdown(outline, source_text, provider)

    print(f"[4/4] 渲染为 docx：{output}")
    render_markdown_to_docx(markdown, template, output)

    total = condense_usage + rewrite_usage
    print("\n完成。")
    print(
        f"Token：input {total.input_tokens:,} / output {total.output_tokens:,} / "
        f"total {total.total_tokens:,}"
    )
    print(f"改写后文档：{output.resolve()}")


def _indent(text: str, prefix: str = "      ") -> str:
    return "\n".join(prefix + line for line in text.splitlines())


def main() -> None:
    parser = argparse.ArgumentParser(description="按模板改写源文件并输出 docx")
    parser.add_argument("--source", required=True, type=Path, help="源文件 .docx 路径")
    parser.add_argument("--template", required=True, type=Path, help="模板 .docx 路径")
    parser.add_argument(
        "--output",
        type=Path,
        default=None,
        help="输出 .docx 路径（默认在源文件同目录生成 *_按模板改写.docx）",
    )
    args = parser.parse_args()

    source: Path = args.source
    template: Path = args.template
    if not source.exists():
        raise SystemExit(f"源文件不存在：{source}")
    if not template.exists():
        raise SystemExit(f"模板不存在：{template}")

    output: Path = args.output or source.with_name(f"{source.stem}_按模板改写.docx")
    asyncio.run(run(source, template, output))


if __name__ == "__main__":
    main()
