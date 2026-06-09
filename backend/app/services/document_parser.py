from __future__ import annotations

import csv
import io
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from pypdf import PdfReader

# 非分页格式（Word/CSV/Excel）没有可靠的物理页概念，这里按字符数切「逻辑页」，
# 让追问检索仍能给出有意义的页码引用，并与下游分片逻辑保持一致。
LOGICAL_PAGE_TARGET_CHARS = 3000
# 表格类数据每个数据块携带的最大行数；过大会超出逻辑页粒度，过小会重复表头浪费空间。
TABULAR_ROWS_PER_BLOCK = 50


@dataclass(frozen=True)
class ParsedPage:
    page_number: int
    text: str


@dataclass(frozen=True)
class ParsedDocument:
    file_path: Path
    file_type: str
    page_count: int | None
    text: str
    pages: list[ParsedPage]


def parse_document(file_path: Path) -> ParsedDocument:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(file_path)
    if suffix == ".docx":
        return parse_docx(file_path)
    if suffix == ".csv":
        return parse_csv(file_path)
    if suffix == ".xlsx":
        return parse_xlsx(file_path)
    if suffix == ".doc":
        raise ValueError("暂不支持旧版 .doc 格式，请用 Word 另存为 .docx 后重试。")
    if suffix == ".xls":
        raise ValueError("暂不支持旧版 .xls 格式，请用 Excel 另存为 .xlsx 后重试。")
    raise ValueError(
        f"暂不支持的文件类型：{suffix or '未知'}。"
        "当前支持 PDF、Word(.docx)、CSV、Excel(.xlsx)。"
    )


def parse_pdf(file_path: Path) -> ParsedDocument:
    reader = PdfReader(str(file_path))
    pages: list[ParsedPage] = []

    for index, page in enumerate(reader.pages, start=1):
        pages.append(ParsedPage(page_number=index, text=page.extract_text() or ""))

    text = "\n\n".join(page.text.strip() for page in pages if page.text.strip())

    return ParsedDocument(
        file_path=file_path,
        file_type="pdf",
        page_count=len(reader.pages),
        text=text,
        pages=pages,
    )


def parse_docx(file_path: Path) -> ParsedDocument:
    document = Document(str(file_path))
    blocks = list(_iter_docx_blocks(document))
    return _build_document(file_path, "docx", blocks)


def parse_csv(file_path: Path) -> ParsedDocument:
    raw = _read_text_with_fallback(file_path)
    rows = [row for row in csv.reader(io.StringIO(raw))]
    blocks = _build_tabular_blocks(rows)
    return _build_document(file_path, "csv", blocks)


def parse_xlsx(file_path: Path) -> ParsedDocument:
    workbook = load_workbook(str(file_path), read_only=True, data_only=True)
    blocks: list[str] = []
    try:
        for worksheet in workbook.worksheets:
            rows = [
                [_cell_to_str(value) for value in row]
                for row in worksheet.iter_rows(values_only=True)
                if any(value is not None for value in row)
            ]
            blocks.extend(
                _build_tabular_blocks(rows, sheet_label=f"工作表：{worksheet.title}")
            )
    finally:
        workbook.close()
    return _build_document(file_path, "xlsx", blocks)


def _build_document(
    file_path: Path,
    file_type: str,
    blocks: Iterable[str],
) -> ParsedDocument:
    pages = _paginate(blocks)
    text = "\n\n".join(page.text for page in pages)
    return ParsedDocument(
        file_path=file_path,
        file_type=file_type,
        page_count=len(pages),
        text=text,
        pages=pages,
    )


def _paginate(
    blocks: Iterable[str],
    target_chars: int = LOGICAL_PAGE_TARGET_CHARS,
) -> list[ParsedPage]:
    """把若干文本块按字符预算打包成逻辑页，单块不再二次切分以保持表格/段落完整。"""
    pages: list[ParsedPage] = []
    current: list[str] = []
    current_len = 0
    page_no = 1
    for block in blocks:
        text = block.strip()
        if not text:
            continue
        if current and current_len + len(text) > target_chars:
            pages.append(ParsedPage(page_number=page_no, text="\n\n".join(current)))
            page_no += 1
            current = []
            current_len = 0
        current.append(text)
        current_len += len(text)
    if current:
        pages.append(ParsedPage(page_number=page_no, text="\n\n".join(current)))
    return pages


def _iter_docx_blocks(document: Document) -> Iterable[str]:
    """按文档原始顺序产出段落与表格文本块（段落与表格交错时也保持顺序）。"""
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            text = Paragraph(child, document).text.strip()
            if text:
                yield text
        elif tag.endswith("}tbl"):
            table_text = _docx_table_to_text(Table(child, document))
            if table_text:
                yield table_text


def _docx_table_to_text(table: Table) -> str:
    lines: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if any(cells):
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _build_tabular_blocks(
    rows: list[list[str]],
    *,
    sheet_label: str | None = None,
    rows_per_block: int = TABULAR_ROWS_PER_BLOCK,
) -> list[str]:
    """把表格行切成块，每块带上工作表标签和列头，便于模型理解每段数据的含义。"""
    cleaned = [[_cell_to_str(cell) for cell in row] for row in rows]
    cleaned = [row for row in cleaned if any(value.strip() for value in row)]
    if not cleaned:
        return []

    column_header = cleaned[0]
    data_rows = cleaned[1:]
    header_line = _join_row(column_header)

    if not data_rows:
        prefix = [sheet_label] if sheet_label else []
        return ["\n".join([*prefix, header_line])]

    blocks: list[str] = []
    for start in range(0, len(data_rows), rows_per_block):
        chunk = data_rows[start : start + rows_per_block]
        lines: list[str] = []
        if sheet_label:
            lines.append(sheet_label)
        lines.append(header_line)
        lines.extend(_join_row(row) for row in chunk)
        blocks.append("\n".join(lines))
    return blocks


def _join_row(row: list[str]) -> str:
    return " | ".join(row)


def _cell_to_str(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _read_text_with_fallback(file_path: Path) -> str:
    """CSV 常见编码兜底：UTF-8(含 BOM) -> GBK -> Latin-1，避免中文导出乱码或解码失败。"""
    raw = file_path.read_bytes()
    for encoding in ("utf-8-sig", "gbk"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("latin-1")
