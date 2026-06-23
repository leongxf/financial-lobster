"""模板 docx 解析与 Markdown 渲染（从 scripts/rewrite_by_template.py 迁入）。"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def _style_heading_level(paragraph: Paragraph) -> int | None:
    """从段落「样式名」推断标题级别；非标题样式返回 None。"""
    style = paragraph.style
    name = (style.name if style is not None else "") or ""
    lowered = name.lower()
    if lowered.startswith("heading") or name.startswith("标题"):
        match = re.search(r"(\d+)", name)
        return int(match.group(1)) if match else 1
    if lowered == "title" or name == "题目":
        return 1
    return None


def _para_signals(paragraph: Paragraph, child) -> tuple[float | None, bool, bool]:
    sizes = [r.font.size.pt for r in paragraph.runs if r.font.size is not None]
    max_size = max(sizes) if sizes else None
    visible_runs = [r for r in paragraph.runs if r.text.strip()]
    any_bold = any(bool(r.bold) for r in visible_runs)
    pPr = child.find(qn("w:pPr"))
    numbered = pPr is not None and pPr.find(qn("w:numPr")) is not None
    return max_size, any_bold, numbered


def _infer_level(paragraph: Paragraph, child) -> int | None:
    level = _style_heading_level(paragraph)
    if level is not None:
        return level
    max_size, any_bold, _ = _para_signals(paragraph, child)
    if max_size is not None and max_size >= 14:
        return 1
    if any_bold:
        return 2
    return None


def extract_template_outline(template_path: Path) -> str:
    """把模板 docx 序列化成带层级的大纲文本。"""
    document = Document(str(template_path))
    body = document.element.body
    lines: list[str] = []
    for child in body.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if not text:
                continue
            level = _infer_level(paragraph, child)
            if level is not None:
                lines.append("\n" + "#" * min(level, 6) + " " + text)
            else:
                lines.append(f"> 填写要求：{text}")
        elif tag.endswith("}tbl"):
            lines.append("> （此处模板含一个表格，若有对应数据请填为 Markdown 表格）")
    outline = "\n".join(lines).strip()
    return outline or "（模板未检测到明确的标题结构，请按源文件内容组织成一篇通顺成稿）"


def parse_template_nodes(template_path: Path, grounded_chapter_keyword: str = "行业分析") -> list[dict]:
    """把模板拆成有序节点：每个标题 + 其下填写要求 + 是否需联网。"""
    document = Document(str(template_path))
    body = document.element.body
    blocks: list[tuple[str, int | None]] = []
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if text:
                blocks.append((text, _infer_level(paragraph, child)))

    nodes: list[dict] = []
    chapter = ""
    i = 0
    n = len(blocks)
    while i < n:
        text, level = blocks[i]
        if level is None:
            i += 1
            continue
        reqs: list[str] = []
        j = i + 1
        while j < n and blocks[j][1] is None:
            reqs.append(blocks[j][0])
            j += 1
        if level == 1:
            chapter = text
        grounded = grounded_chapter_keyword in chapter and level >= 2 and bool(reqs)
        nodes.append({"level": level, "title": text, "reqs": reqs, "grounded": grounded})
        i = j
    return nodes


def group_nodes_by_chapter(nodes: list[dict]) -> list[dict]:
    """把扁平节点按 L1 章分组，每章含 subsections（L2+ 直到下一 L1）。"""
    chapters: list[dict] = []
    current: dict | None = None
    for node in nodes:
        if node["level"] == 1:
            if current is not None:
                chapters.append(current)
            current = {
                "title": node["title"],
                "level": 1,
                "reqs": node["reqs"],
                "subsections": [],
            }
        elif current is not None and node["level"] >= 2:
            current["subsections"].append(node)
    if current is not None:
        chapters.append(current)
    return chapters


def resolve_web_fill_sections(nodes: list[dict], chapter_keywords: list[str]) -> list[dict]:
    """按配置的章关键词，解析需联网填充的小节列表。

    匹配到某 L1 章后：若该章有子标题（L2+），则对每个子标题分别检索填充；
    若无子标题，则填充该章本身。
    """
    if not chapter_keywords:
        return []

    sections: list[dict] = []
    for chapter in group_nodes_by_chapter(nodes):
        if not any(kw in chapter["title"] for kw in chapter_keywords):
            continue
        chapter_reqs = chapter["reqs"]
        chapter_title = chapter["title"]
        if chapter["subsections"]:
            for sub in chapter["subsections"]:
                sections.append(
                    {
                        **sub,
                        "chapter_title": chapter_title,
                        "chapter_reqs": chapter_reqs,
                    }
                )
        else:
            sections.append(
                {
                    "level": 1,
                    "title": chapter_title,
                    "reqs": chapter_reqs,
                    "chapter_title": chapter_title,
                    "chapter_reqs": [],
                }
            )
    return sections


def effective_section_reqs(section: dict) -> list[str]:
    """合并章级说明与子节填写要求，供检索词生成与正文撰写使用。"""
    reqs: list[str] = []
    chapter_title = section.get("chapter_title")
    if chapter_title:
        reqs.append(f"所属章节：{chapter_title}")
    reqs.extend(section.get("chapter_reqs") or [])
    reqs.extend(section.get("reqs") or [])
    return reqs


def _clear_body(document: Document) -> None:
    body = document.element.body
    for child in list(body.iterchildren()):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _add_runs_with_bold(paragraph: Paragraph, text: str) -> None:
    for index, segment in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not segment:
            continue
        run = paragraph.add_run(segment)
        if index % 2 == 1:
            run.bold = True


def _paragraph_safe(document: Document, style_candidates: list[str]):
    for style in style_candidates:
        try:
            return document.add_paragraph(style=style), style
        except KeyError:
            continue
    return document.add_paragraph(), None


def _add_heading_safe(document: Document, text: str, level: int):
    try:
        return document.add_heading(text, level=min(level, 9))
    except KeyError:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        return paragraph


def _add_list_item(document: Document, text: str, *, ordered: bool, index: int) -> None:
    candidates = ["List Number"] if ordered else ["List Bullet"]
    paragraph, used = _paragraph_safe(document, candidates)
    if used is None:
        paragraph, _ = _paragraph_safe(document, ["List Paragraph"])
        prefix = f"{index}. " if ordered else "• "
        _add_runs_with_bold(paragraph, prefix + text)
        return
    _add_runs_with_bold(paragraph, text)


def _is_table_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?[\s:\-\|]+\|?\s*", line)) and "-" in line


def _split_table_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def _render_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = row[c_idx] if c_idx < len(row) else ""
            if r_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def render_markdown_to_docx(markdown: str, template_path: Path, output_path: Path) -> None:
    """以模板为基底渲染 Markdown 成稿为 docx（继承模板主题/页眉页脚）。"""
    document = Document(str(template_path))
    _clear_body(document)

    lines = markdown.replace("\r\n", "\n").split("\n")
    i = 0
    n = len(lines)
    ordered_index = 0
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            ordered_index = 0
            table_lines: list[str] = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [
                _split_table_row(tl) for tl in table_lines if not _is_table_separator(tl)
            ]
            _render_table(document, rows)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            ordered_index = 0
            level = len(heading_match.group(1))
            _add_heading_safe(document, heading_match.group(2).strip(), level)
            i += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            ordered_index = 0
            _add_list_item(document, bullet_match.group(1).strip(), ordered=False, index=0)
            i += 1
            continue

        ordered_match = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if ordered_match:
            ordered_index += 1
            _add_list_item(
                document, ordered_match.group(1).strip(), ordered=True, index=ordered_index
            )
            i += 1
            continue

        ordered_index = 0
        paragraph = document.add_paragraph()
        _add_runs_with_bold(paragraph, stripped)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
